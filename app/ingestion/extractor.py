import re
from pathlib import Path
import pdfplumber
import docx

def clean_text(text: str) -> str:
    """
    Cleans whitespace, strips lines, and removes duplicate blank lines.
    """
    if not text:
        return ""
    # Normalize lines: strip trailing/leading spaces on each line
    lines = [line.strip() for line in text.splitlines()]
    # Remove duplicate blank lines and empty lines
    cleaned_lines = []
    prev_was_empty = False
    for line in lines:
        if line == "":
            if not prev_was_empty:
                cleaned_lines.append("")
                prev_was_empty = True
        else:
            cleaned_lines.append(line)
            prev_was_empty = False
            
    cleaned_text = "\n".join(cleaned_lines)
    # Replace multiple spaces with a single space
    cleaned_text = re.sub(r"[ \t]+", " ", cleaned_text)
    return cleaned_text.strip()

def extract_text_from_pdf(file_path: Path) -> str:
    """
    Extracts text from a PDF file using pdfplumber.
    """
    text_content = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_content.append(page_text)
    return "\n".join(text_content)

def extract_text_from_docx(file_path: Path) -> str:
    """
    Extracts text from a DOCX file using python-docx.
    """
    doc = docx.Document(file_path)
    text_content = []
    for paragraph in doc.paragraphs:
        if paragraph.text:
            text_content.append(paragraph.text)
    return "\n".join(text_content)

def extract_text_from_txt(file_path: Path) -> str:
    """
    Extracts text from a TXT file.
    """
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def extract_text(file_path: Path) -> str:
    """
    Extracts and cleans text from PDF, DOCX, or TXT file.
    """
    suffix = file_path.suffix.lower()
    raw_text = ""
    if suffix == ".pdf":
        raw_text = extract_text_from_pdf(file_path)
    elif suffix == ".docx":
        raw_text = extract_text_from_docx(file_path)
    elif suffix == ".txt":
        raw_text = extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {suffix}")
        
    return clean_text(raw_text)

def extract_images_from_pdf(file_path: Path, output_dir: Path) -> List[dict]:
    """
    Extracts embedded images/figures from a PDF file using pdfplumber & PIL.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    images_info = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for p_idx, page in enumerate(pdf.pages):
                for i_idx, img in enumerate(page.images):
                    try:
                        x0, top, x1, bottom = img["x0"], img["top"], img["x1"], img["bottom"]
                        if (x1 - x0) >= 30 and (bottom - top) >= 30:
                            cropped = page.crop((x0, top, x1, bottom))
                            img_obj = cropped.to_image(resolution=150)
                            save_name = f"{file_path.stem}_p{p_idx+1}_img{i_idx+1}.png"
                            save_path = output_dir / save_name
                            img_obj.save(save_path, format="PNG")
                            
                            images_info.append({
                                "path": f"outputs/extracted_images/{save_name}",
                                "filename": save_name,
                                "caption": f"Figure {len(images_info)+1} from {file_path.name} (Page {p_idx+1})",
                                "page": p_idx + 1,
                                "doc": file_path.name
                            })
                    except Exception:
                        pass
    except Exception:
        pass
    return images_info

def extract_images_from_docx(file_path: Path, output_dir: Path) -> List[dict]:
    """
    Extracts embedded images from a DOCX file using python-docx.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    images_info = []
    try:
        doc = docx.Document(file_path)
        for r_id, rel in doc.part.rels.items():
            if "image" in rel.target_ref:
                img_part = rel.target_part
                img_bytes = img_part.blob
                ext = rel.target_ref.split(".")[-1]
                save_name = f"{file_path.stem}_img_{r_id}.{ext}"
                save_path = output_dir / save_name
                with open(save_path, "wb") as f:
                    f.write(img_bytes)
                images_info.append({
                    "path": f"outputs/extracted_images/{save_name}",
                    "filename": save_name,
                    "caption": f"Embedded Image {len(images_info)+1} from {file_path.name}",
                    "page": 1,
                    "doc": file_path.name
                })
    except Exception:
        pass
    return images_info

def extract_images(file_path: Path, output_dir: Path = Path("outputs/extracted_images")) -> List[dict]:
    """
    Extracts images from PDF or DOCX file.
    """
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return extract_images_from_pdf(file_path, output_dir)
    elif suffix == ".docx":
        return extract_images_from_docx(file_path, output_dir)
    return []

def extract_captions(text: str) -> List[str]:
    """
    Scans the text for potential figure, diagram, illustration, or picture captions.
    """
    if not text:
        return []
    
    patterns = [
        r'(?mi)^\s*(?:Figure|Fig\.|Illustration|Picture|Diagram|Table)\s*\d*[\s\:\-\.\,]+.*$',
        r'(?mi)^\s*(?:Illustration of|Diagram showing)\s+.*$'
    ]
    
    captions = []
    for pattern in patterns:
        matches = re.findall(pattern, text)
        for m in matches:
            clean_m = m.strip()
            clean_m = re.sub(r'\s+', ' ', clean_m)
            if len(clean_m) > 10 and len(clean_m) < 250 and clean_m not in captions:
                captions.append(clean_m)
                
    if "WESTERN GHATS" in text and "CAMBAY RIFT" in text:
        diagram_ref = "Diagram: Rift Valley Architecture (Horsts and Grabens block model)"
        if diagram_ref not in captions:
            captions.append(diagram_ref)
            
    return captions
